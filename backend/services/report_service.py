"""报告导出服务（毕设功能）。"""

from __future__ import annotations

import hashlib
import logging
import os
from collections.abc import Callable
from functools import lru_cache
from html import escape
from pathlib import Path
from typing import Any

from PIL import Image as PILImage

from backend.utils.cache import TTLCache
from backend.utils.path_utils import validate_outcrop_name
from trace_pipeline.config import PROJECT_ROOT
from trace_pipeline.geology.statistics import TraceStatisticsConfig, compute_trace_statistics
from trace_pipeline.pipeline import load_trace_data
from trace_pipeline.utils.fonts import is_cjk

logger = logging.getLogger(__name__)

REPORT_DIR = PROJECT_ROOT / "reports"
_REPORT_CACHE_TTL = 300.0

_LATIN_FONT = "TimesNewRoman"
_LATIN_FALLBACK_FONT = "Times-Roman"
_CJK_BODY_FONT = "ReportBodyCJK"
_CJK_HEADING_FONT = "ReportHeadingCJK"


def _font_candidates(kind: str) -> list[tuple[str, str]]:
    windir = os.environ.get("WINDIR", r"C:\Windows")
    fonts_dir = os.path.join(windir, "Fonts")
    user_fonts = os.path.expanduser("~/.fonts")
    if kind == "latin":
        return [
            (os.path.join(fonts_dir, "times.ttf"), "Times New Roman"),
            (os.path.join(fonts_dir, "timesbd.ttf"), "Times New Roman Bold"),
            (
                "/usr/share/fonts/truetype/msttcorefonts/Times_New_Roman.ttf",
                "Times New Roman",
            ),
            (
                "/usr/share/fonts/truetype/liberation/LiberationSerif-Regular.ttf",
                "Liberation Serif",
            ),
            ("/usr/share/fonts/truetype/dejavu/DejaVuSerif.ttf", "DejaVu Serif"),
        ]
    if kind == "heading":
        return [
            (os.path.join(fonts_dir, "simhei.ttf"), "SimHei"),
            (os.path.join(fonts_dir, "msyhbd.ttc"), "Microsoft YaHei Bold"),
            (os.path.join(fonts_dir, "msyh.ttc"), "Microsoft YaHei"),
            ("/System/Library/Fonts/STHeiti Light.ttc", "STHeiti"),
            ("/System/Library/Fonts/PingFang.ttc", "PingFang"),
            ("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc", "Noto Sans CJK"),
            ("/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc", "WenQuanYi Zen Hei"),
            (os.path.join(user_fonts, "NotoSansCJK-Regular.ttc"), "Noto Sans CJK"),
        ]
    return [
        (os.path.join(fonts_dir, "simsun.ttc"), "SimSun"),
        (os.path.join(fonts_dir, "SimSun.ttf"), "SimSun"),
        ("/System/Library/Fonts/Supplemental/Songti.ttc", "Songti"),
        ("/System/Library/Fonts/STSong.ttf", "STSong"),
        ("/usr/share/fonts/opentype/noto/NotoSerifCJK-Regular.ttc", "Noto Serif CJK"),
        ("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc", "Noto Sans CJK"),
        ("/usr/share/fonts/truetype/wqy/wqy-microhei.ttc", "WenQuanYi Micro Hei"),
        (os.path.join(user_fonts, "NotoSerifCJK-Regular.ttc"), "Noto Serif CJK"),
        (os.path.join(user_fonts, "NotoSansCJK-Regular.ttc"), "Noto Sans CJK"),
    ]


def _register_pdf_font(pdfmetrics, ttfont_cls, font_name: str, candidates, fallback: str) -> str:
    """注册 ReportLab 字体，成功返回注册名，失败返回 fallback。"""
    for font_path, label in candidates:
        if not font_path or not os.path.exists(font_path):
            continue
        try:
            pdfmetrics.registerFont(ttfont_cls(font_name, font_path))
            return font_name
        except (OSError, ValueError, RuntimeError) as exc:
            logger.warning("注册字体 %s (%s) 失败: %s", label, font_path, exc)
    return fallback


def _pdf_mixed_font_markup(text: str, *, cjk_font: str, latin_font: str) -> str:
    """将中英文混排文本拆成 ReportLab font 片段。"""
    blocks: list[str] = []
    current_font = ""
    current_text = ""
    for ch in str(text):
        font = cjk_font if is_cjk(ch) else latin_font
        if not current_font:
            current_font = font
            current_text = ch
        elif font == current_font:
            current_text += ch
        else:
            blocks.append(f'<font name="{current_font}">{escape(current_text)}</font>')
            current_font = font
            current_text = ch
    if current_text:
        blocks.append(f'<font name="{current_font}">{escape(current_text)}</font>')
    return "".join(blocks)


@lru_cache(maxsize=1)
def _find_system_font() -> tuple[str, str]:
    """跨平台字体探测：返回 (font_path, font_name)。"""
    import platform

    system = platform.system()

    if system == "Windows":
        windir = os.environ.get("WINDIR", r"C:\Windows")
        fonts_dir = os.path.join(windir, "Fonts")
        windows_candidates = [
            (os.path.join(fonts_dir, "simsun.ttc"), "SimSun"),
            (os.path.join(fonts_dir, "SimSun.ttf"), "SimSun"),
            (os.path.join(fonts_dir, "msyh.ttc"), "MicrosoftYaHei"),
            (os.path.join(fonts_dir, "msyhbd.ttc"), "MicrosoftYaHei"),
            (os.path.join(fonts_dir, "simhei.ttf"), "SimHei"),
        ]
        candidates = windows_candidates
    else:
        unix_candidates = [
            ("/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc", "WenQuanYiZenHei"),
            ("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", "DejaVuSans"),
            ("/usr/local/share/fonts/truetype/wqy/wqy-zenhei.ttc", "WenQuanYiZenHei"),
            ("/System/Library/Fonts/PingFang.ttc", "PingFang"),
            ("/System/Library/Fonts/STHeiti Light.ttc", "STHeiti"),
            ("/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf", "LiberationSans"),
        ]
        user_fonts = os.path.expanduser("~/.fonts")
        for name, label in [
            ("wqy-zenhei.ttc", "WenQuanYiZenHei"),
            ("NotoSansCJK-Regular.ttc", "NotoSansCJK"),
        ]:
            unix_candidates.append((os.path.join(user_fonts, name), label))
        candidates = unix_candidates

    for fp, fname in candidates:
        if os.path.exists(fp):
            return fp, fname

    # 回退：尝试从 matplotlib 字体管理器查找 CJK 字体
    try:
        import matplotlib.font_manager as fm

        cjk_keywords = [
            "noto",
            "cjk",
            "wqy",
            "sourcehansans",
            "sourcehanserifs",
            "microsoftyahei",
            "simsun",
            "simhei",
            "pingfang",
            "heiti",
            "meiryo",
            "malgun",
            "yugothic",
            "nanum",
        ]
        for font in fm.fontManager.ttflist:
            name_lower = font.name.lower()
            if any(k in name_lower for k in cjk_keywords):
                return font.fname, font.name
        # 回退到系统默认 sans-serif 字体（通常能渲染拉丁字符）
        default_font = fm.findfont(fm.FontProperties(family="sans-serif"))
        if default_font and os.path.exists(default_font):
            return default_font, "sans-serif"
    except Exception as exc:
        logger.debug("CJK 字体检测失败: %s", exc)

    return "", ""


class ReportService:
    """生成 Word / PDF 报告。"""

    def __init__(self) -> None:
        # 缓存已生成的报告结果，避免同一配置下重复写入 DOCX/PDF
        self._result_cache = TTLCache(ttl=_REPORT_CACHE_TTL)

    @staticmethod
    def _cache_key(outcrop: str, report_type: str, fmt: str, config: dict[str, Any]) -> str:
        # 取影响报告内容的配置子集
        cfg_part = (
            config.get("input_dir", ""),
            config.get("output_dir", ""),
            config.get("window_strategy", "auto"),
            config.get("min_intersections", 5),
            config.get("rose_bin_width", 10.0),
        )
        # 使用稳定哈希（sha256）避免 PYTHONHASHSEED 随机化导致进程重启后缓存全部失效
        cfg_hash = hashlib.sha256(repr(cfg_part).encode()).hexdigest()[:16]
        return f"{outcrop}:{report_type}:{fmt}:{cfg_hash}"

    @staticmethod
    def _image_mtimes(img_paths: list[str]) -> dict[str, float]:
        mtimes: dict[str, float] = {}
        for p in img_paths:
            try:
                mtimes[p] = Path(p).stat().st_mtime
            except OSError:
                mtimes[p] = 0.0
        return mtimes

    def _try_cached(
        self, outcrop: str, report_type: str, fmt: str, config: dict[str, Any]
    ) -> dict[str, Any] | None:
        key = self._cache_key(outcrop, report_type, fmt, config)
        cached = self._result_cache.get(key)
        if cached is None:
            return None
        result, img_mtimes = cached["result"], cached["img_mtimes"]
        # 若报告引用的输出图片在缓存后被修改/删除，则失效缓存
        for path, mtime in img_mtimes.items():
            try:
                if Path(path).stat().st_mtime != mtime:
                    return None
            except OSError:
                return None
        return result

    def _store_cached(
        self,
        outcrop: str,
        report_type: str,
        fmt: str,
        config: dict[str, Any],
        result: dict[str, Any],
        img_paths: list[str],
    ) -> None:
        key = self._cache_key(outcrop, report_type, fmt, config)
        self._result_cache.set(
            key, {"result": result, "img_mtimes": self._image_mtimes(img_paths)}
        )

    def generate(
        self,
        outcrop: str,
        report_type: str,
        fmt: str,
        config: dict[str, Any],
        progress_callback: Callable[[str, str], None] | None = None,
    ) -> dict[str, Any]:
        """生成报告并返回文件路径。

        progress_callback(step, message) 用于向调用方报告进度。step 取值:
          "loading", "stats", "docx", "pdf", "done"
        """
        try:
            validate_outcrop_name(outcrop)
        except ValueError as exc:
            return {"error": str(exc)}
        import time

        start = time.perf_counter()

        REPORT_DIR.mkdir(parents=True, exist_ok=True)
        results = {}

        input_dir = config.get("input_dir", "input")
        table_stem = f"{outcrop}_process"
        logger.info(
            "报告生成开始 [%s]: type=%s, fmt=%s",
            outcrop,
            report_type,
            fmt,
            extra={
                "stage": "report_start",
                "outcrop": outcrop,
                "report_type": report_type,
                "fmt": fmt,
            },
        )
        try:
            if progress_callback:
                progress_callback("loading", "正在加载迹线数据...")
            trace = load_trace_data(input_dir, table_stem, outcrop)
            stats_config = TraceStatisticsConfig(
                window_strategy=config.get("window_strategy", "auto"),
                min_intersections=config.get("min_intersections", 5),
            )
            if progress_callback:
                progress_callback("stats", "正在计算统计量...")
            statistics = compute_trace_statistics(trace, stats_config)
        except Exception as exc:
            logger.warning(
                "报告 [%s] 数据加载失败: %s",
                outcrop,
                exc,
                extra={"stage": "report_error", "outcrop": outcrop, "error": str(exc)},
            )
            return {"error": str(exc)}

        ctx = self._build_report_context(outcrop, trace, statistics, report_type, config)
        cached = self._try_cached(outcrop, report_type, fmt, config)
        if cached is not None:
            logger.info(
                "报告 [%s] 命中缓存，直接返回已有结果",
                outcrop,
                extra={"stage": "report_cache_hit", "outcrop": outcrop},
            )
            if progress_callback:
                progress_callback("done", "命中缓存，直接返回")
            return cached

        if fmt in ("docx", "both"):
            if progress_callback:
                progress_callback("docx", "正在生成 DOCX...")
            docx_result = self._gen_docx(outcrop, ctx)
            if "error" in docx_result:
                results["docx_error"] = docx_result["error"]
            else:
                results["docx"] = docx_result["path"]
                logger.info(
                    "DOCX 报告生成: %s",
                    docx_result["path"],
                    extra={"stage": "report_docx", "outcrop": outcrop, "path": docx_result["path"]},
                )
        if fmt in ("pdf", "both"):
            if progress_callback:
                progress_callback("pdf", "正在生成 PDF...")
            pdf_result = self._gen_pdf(outcrop, ctx)
            if "error" in pdf_result:
                results["pdf_error"] = pdf_result["error"]
            else:
                results["pdf"] = pdf_result["path"]
                logger.info(
                    "PDF 报告生成: %s",
                    pdf_result["path"],
                    extra={"stage": "report_pdf", "outcrop": outcrop, "path": pdf_result["path"]},
                )

        self._store_cached(outcrop, report_type, fmt, config, results, ctx.get("img_paths", []))

        if progress_callback:
            progress_callback("done", "报告生成完毕")

        duration = (time.perf_counter() - start) * 1000
        logger.info(
            "报告生成完成 [%s]: docx=%s, pdf=%s (%.3f ms)",
            outcrop,
            bool(results.get("docx")),
            bool(results.get("pdf")),
            duration,
            extra={
                "stage": "report_complete",
                "outcrop": outcrop,
                "report_type": report_type,
                "fmt": fmt,
                "has_docx": bool(results.get("docx")),
                "has_pdf": bool(results.get("pdf")),
                "duration_ms": round(duration, 3),
            },
        )
        return results

    # ------------------------------------------------------------------
    # 共享：提取报告数据上下文（消除 docx/pdf 之间的重复）
    # ------------------------------------------------------------------
    def _build_report_context(
        self, outcrop: str, trace, statistics, report_type: str, config: dict[str, Any]
    ) -> dict[str, Any]:
        """返回报告所需的统计文本行和图片路径，供各格式渲染器共用。"""
        stat_lines: list[str] = []
        if report_type in ("full", "stats"):
            stat_lines = [
                f"露头标识: {outcrop}",
                f"测线走向: {trace.scanline_azimuth:.2f}°",
                f"迹线条数: {trace.count}",
                f"平均迹长: {statistics.mean_trace_length:.4f} m",
                f"线密度 P10: {statistics.p10:.4f} m⁻¹",
                f"面密度 P20: {statistics.p20:.4f} m⁻²",
                f"长度密度 P21: {statistics.p21:.4f} m⁻¹",
                f"测线长度: {statistics.scanline_length:.4f} m",
                (
                    f"露头面积: {statistics.outcrop_area:.4f} m² "
                    f"(来源: {statistics.outcrop_area_source or 'unknown'})"
                ),
                f"圆窗策略: {statistics.window_strategy or 'auto'}",
                f"I 型迹线数: {statistics.type_i_count}",
                f"II 型迹线数: {statistics.type_ii_count}",
                f"III 型迹线数: {statistics.type_iii_count}",
            ]
            if statistics.window_validation_warning:
                stat_lines.append(f"警告: {statistics.window_validation_warning}")

        img_names: list[str] = []
        if report_type in ("full", "plots"):
            out_dir = Path(config.get("output_dir", "output"))
            if not out_dir.is_absolute():
                out_dir = PROJECT_ROOT / out_dir
            out_dir = out_dir.resolve()
            for img_name in [
                f"{outcrop}_raw(n={trace.count}).png",
                f"{outcrop}_rotated(strike={trace.scanline_azimuth:.1f}).png",
                f"{outcrop}_rose(bin={config.get('rose_bin_width', 10.0)}).png",
            ]:
                img_path = out_dir / img_name
                if img_path.exists():
                    img_names.append(str(img_path))

        return {
            "title": f"{outcrop} 迹线分析报告",
            "stat_lines": stat_lines,
            "img_paths": img_names,
        }

    # ------------------------------------------------------------------
    # Word
    # ------------------------------------------------------------------
    def _gen_docx(self, outcrop: str, ctx: dict[str, Any]) -> dict[str, str]:
        try:
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn
            from docx.shared import Inches, Pt
        except ImportError:
            logger.warning("python-docx 未安装")
            return {"error": "python-docx 未安装"}

        try:
            doc = Document()

            def _set_style_font(
                style, western="Times New Roman", east_asia="SimSun", size=None, bold=False
            ):
                style.font.name = western
                if size is None:
                    size = Pt(12)
                style.font.size = size
                style.font.bold = bold
                if style.element.rPr is not None:
                    style.element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)

            _set_style_font(doc.styles["Normal"])
            for lvl in [0, 1, 2, 3]:
                try:
                    name = "Title" if lvl == 0 else f"Heading {lvl}"
                    _set_style_font(
                        doc.styles[name],
                        size=Pt([20, 16, 14, 12][lvl]),
                        bold=True,
                        east_asia="SimHei",
                    )
                except KeyError:
                    logger.debug("Word 样式 %s 不存在，跳过字体设置", "Title" if lvl == 0 else f"Heading {lvl}")

            def _add_para(text, style=None):
                p = doc.add_paragraph(text, style=style)
                for run in p.runs:
                    run.font.name = "Times New Roman"
                    if run.element.rPr is not None:
                        run.element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
                return p

            title = doc.add_heading(ctx["title"], level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in title.runs:
                run.font.name = "Times New Roman"
                if run.element.rPr is not None:
                    run.element.rPr.rFonts.set(qn("w:eastAsia"), "SimHei")

            for line in ctx["stat_lines"]:
                _add_para(line)
            for img_path in ctx["img_paths"]:
                doc.add_picture(img_path, width=Inches(5.5))

            path = REPORT_DIR / f"{outcrop}_report.docx"
            doc.save(str(path))
            return {"path": str(path.resolve())}
        except Exception as exc:
            logger.exception("DOCX 生成失败: %s", exc)
            return {"error": f"DOCX 生成失败: {exc}"}

    # ------------------------------------------------------------------
    # PDF
    # ------------------------------------------------------------------
    def _gen_pdf(self, outcrop: str, ctx: dict[str, Any]) -> dict[str, str]:
        try:
            from reportlab.lib.enums import TA_CENTER
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.cidfonts import UnicodeCIDFont
            from reportlab.pdfbase.ttfonts import TTFont
            from reportlab.platypus import Image as RLImage
            from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
        except ImportError:
            logger.warning("reportlab 未安装")
            return {"error": "reportlab 未安装"}

        try:
            fallback_path, _fallback_name = _find_system_font()
            fallback_candidates = [(fallback_path, "System CJK")] if fallback_path else []
            latin_font = _register_pdf_font(
                pdfmetrics,
                TTFont,
                _LATIN_FONT,
                _font_candidates("latin"),
                _LATIN_FALLBACK_FONT,
            )
            body_font = _register_pdf_font(
                pdfmetrics,
                TTFont,
                _CJK_BODY_FONT,
                [*_font_candidates("body"), *fallback_candidates],
                latin_font,
            )
            if body_font == latin_font:
                try:
                    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
                    body_font = "STSong-Light"
                except Exception as exc:
                    logger.warning("注册 PDF 中文兜底字体 STSong-Light 失败: %s", exc)
            heading_font = _register_pdf_font(
                pdfmetrics,
                TTFont,
                _CJK_HEADING_FONT,
                [*_font_candidates("heading"), *fallback_candidates],
                body_font,
            )

            doc = SimpleDocTemplate(str(REPORT_DIR / f"{outcrop}_report.pdf"), pagesize=A4)
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                "CustomTitle",
                parent=styles["Heading1"],
                fontName=heading_font,
                fontSize=18,
                alignment=TA_CENTER,
                spaceAfter=20,
            )
            body_style = ParagraphStyle(
                "CustomBody",
                parent=styles["BodyText"],
                fontName=body_font,
                fontSize=11,
                spaceAfter=8,
            )

            story = [
                Paragraph(
                    _pdf_mixed_font_markup(
                        ctx["title"], cjk_font=heading_font, latin_font=latin_font
                    ),
                    title_style,
                ),
                Spacer(1, 12),
            ]
            for line in ctx["stat_lines"]:
                story.append(
                    Paragraph(
                        _pdf_mixed_font_markup(line, cjk_font=body_font, latin_font=latin_font),
                        body_style,
                    )
                )
            if ctx["stat_lines"]:
                story.append(Spacer(1, 12))
            for img_path in ctx["img_paths"]:
                try:
                    with PILImage.open(img_path) as im:
                        orig_w, orig_h = im.size
                    max_width = 400
                    width = max_width
                    height = max_width * orig_h / orig_w if orig_w else 300
                    story.append(RLImage(img_path, width=width, height=height))
                    story.append(Spacer(1, 12))
                except Exception as img_exc:
                    logger.warning("PDF 跳过图片 %s: %s", img_path, img_exc)

            doc.build(story)
            return {"path": str((REPORT_DIR / f"{outcrop}_report.pdf").resolve())}
        except Exception as exc:
            logger.exception("PDF 生成失败: %s", exc)
            return {"error": f"PDF 生成失败: {exc}"}
